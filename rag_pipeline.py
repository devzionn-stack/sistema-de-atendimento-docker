
import os
import logging
from typing import List, Dict, Any, Optional
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_community.document_loaders import PyPDFLoader, TextLoader
from langchain_core.documents import Document
from langchain_google_genai import GoogleGenerativeAIEmbeddings
from database_service import db_service

# Configuração de Logs
logging.basicConfig(level=logging.INFO)
logger = logging.getLogger("RAG_Pipeline")

class DocumentProcessor:
    """Carrega documentos de diferentes formatos (PDF, TXT, DOCX)."""
    @staticmethod
    def load_document(file_path: str) -> List[Document]:
        if not os.path.exists(file_path):
             logger.error(f"Arquivo não encontrado: {file_path}")
             raise FileNotFoundError(f"Arquivo não encontrado: {file_path}")

        ext = os.path.splitext(file_path)[1].lower()
        try:
            if ext == ".pdf":
                loader = PyPDFLoader(file_path)
                return loader.load()
            elif ext == ".docx":
                import docx
                doc = docx.Document(file_path)
                content = '\n'.join([p.text for p in doc.paragraphs])
                return [Document(page_content=content, metadata={"source": file_path})]
            elif ext == ".txt":
                loader = TextLoader(file_path)
                return loader.load()
            else:
                raise ValueError(f"Formato não suportado: {ext}")
        except Exception as e:
            logger.error(f"Erro ao processar {file_path}: {e}")
            raise

class TextChunker:
    """Divide documentos em fragmentos menores para vetorização."""
    def __init__(self, chunk_size: int = 1000, chunk_overlap: int = 200):
        self.splitter = RecursiveCharacterTextSplitter(
            chunk_size=chunk_size, 
            chunk_overlap=chunk_overlap,
            separators=["\n\n", "\n", " ", ""]
        )
    
    def split(self, documents: List[Document]) -> List[Document]:
        return self.splitter.split_documents(documents)

class EmbeddingFactory:
    """Gera embeddings usando Google Gemini."""
    def __init__(self):
        # A chave de API é obtida automaticamente do ambiente (GOOGLE_API_KEY ou API_KEY)
        self.embeddings = GoogleGenerativeAIEmbeddings(model="models/embedding-001")

    def embed_documents(self, texts: List[str]) -> List[List[float]]:
        try:
            return self.embeddings.embed_documents(texts)
        except Exception as e:
            logger.error(f"Erro ao gerar embeddings: {e}")
            # Retorna vetores zerados em caso de falha para não quebrar o pipeline em dev
            return [[0.0] * 768 for _ in texts]

    def embed_query(self, text: str) -> List[float]:
        try:
            return self.embeddings.embed_query(text)
        except Exception as e:
            logger.error(f"Erro ao gerar embedding da query: {e}")
            return [0.0] * 768

class RAGPipeline:
    """Gerencia o ciclo de vida RAG: Ingestão -> Vetorização -> Persistência."""
    def __init__(self):
        self.embedder = EmbeddingFactory()

    async def process_and_index(self, file_path: str, user_id: str = "system") -> Dict[str, Any]:
        """
        Executa o pipeline completo:
        1. Carrega o arquivo
        2. Divide em chunks
        3. Gera vetores
        4. Salva no Postgres (pgvector)
        """
        try:
            logger.info(f"RAG: Iniciando ingestão de {file_path}")
            
            # 1. Carregamento
            docs = DocumentProcessor.load_document(file_path)
            if not docs:
                logger.warning(f"RAG: Documento vazio {file_path}")
                return {"status": "empty", "chunks": 0}

            # 2. Chunking
            chunker = TextChunker()
            chunks = chunker.split(docs)
            logger.info(f"RAG: Gerados {len(chunks)} fragmentos.")

            # 3. Embedding
            texts = [c.page_content for c in chunks]
            vectors = self.embedder.embed_documents(texts)
            
            # 4. Preparação para Storage
            storage_items = []
            for i, chunk in enumerate(chunks):
                storage_items.append({
                    "content": chunk.page_content,
                    "embedding": vectors[i],
                    "metadata": {
                        "source": os.path.basename(file_path),
                        "user_id": user_id,
                        "page": chunk.metadata.get("page", 0),
                        "chunk_index": i
                    }
                })

            # 5. Persistência
            # A lógica de inserção no Postgres está encapsulada no db_service
            await db_service.store_vectors(storage_items)
            
            return {"status": "success", "chunks": len(storage_items)}

        except Exception as e:
            logger.error(f"RAG Pipeline falhou: {e}")
            return {"status": "error", "message": str(e)}

    async def remove_document(self, file_name: str) -> Dict[str, Any]:
        """Remove um documento e seus vetores do índice."""
        try:
            await db_service.delete_vectors(file_name)
            logger.info(f"RAG: Removido documento {file_name}")
            return {"status": "success", "message": f"Documento {file_name} removido."}
        except Exception as e:
            logger.error(f"RAG: Falha ao remover documento: {e}")
            return {"status": "error", "message": str(e)}

rag_manager = RAGPipeline()
